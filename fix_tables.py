import docx, sys
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH


if __name__ == "__main__":


    input_fname = sys.argv[1]
    out_fname = sys.argv[2]

    document = docx.Document(input_fname) # read in the doc file

    print("Fixing table alignment...")

    # grab each table in the document
    for table in document.tables:

        # align to center
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.allow_autofit = True

        for row in table.rows:
            for cell in row.cells:

                # align each cell's contents to the center
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    document.save(out_fname) # save out